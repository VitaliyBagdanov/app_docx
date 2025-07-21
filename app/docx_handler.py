# import os
import re
from docx import Document

def extract_tags_from_docx(template_path: str):
    """
    Извлекает все плейсхолдеры вида {tag} из docx-шаблона.
    """
    doc = Document(template_path)
    tags = set()
    pattern = re.compile(r'{(\w+)}')

    # Поиск в параграфах
    for paragraph in doc.paragraphs:
        tags.update(pattern.findall(paragraph.text))

    # Поиск в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                tags.update(pattern.findall(cell.text))

    return list(tags)

def replace_placeholder_in_paragraph(paragraph, values: dict):
    """
    Корректно заменяет плейсхолдеры в параграфе, включая работу с runs.
    """
    pattern = re.compile(r'{(\w+)}')
    for key, val in values.items():
        placeholder = f"{{{key}}}"
        if placeholder in paragraph.text:
            # Для "разорванных" runs
            inline = paragraph.runs
            for i in range(len(inline)):
                if placeholder in inline[i].text:
                    inline[i].text = inline[i].text.replace(placeholder, str(val))
            # Если всё единым run
            paragraph.text = paragraph.text.replace(placeholder, str(val))

def replace_placeholders_in_table(table, values: dict):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_placeholder_in_paragraph(paragraph, values)

def generate_filled_docx(template_path: str, output_path: str, values: dict):
    """
    Заменяет плейсхолдеры вида {key} на значения из values и сохраняет в output_path.
    """
    doc = Document(template_path)

    # Заменяем во всех параграфах
    for paragraph in doc.paragraphs:
        replace_placeholder_in_paragraph(paragraph, values)

    # Заменяем внутри таблиц
    for table in doc.tables:
        replace_placeholders_in_table(table, values)

    doc.save(output_path)
